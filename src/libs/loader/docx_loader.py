"""DOCX Loader with semantic-aware parsing.

This module implements DOCX parsing with preservation of:
- Document structure (headings, lists, tables)
- Styles and formatting (bold, italic, colors)
- Comments/annotations (批注)
- Track changes/revisions (修订)
- Embedded images
- Headers and footers

Key Features:
- Preserves "hidden knowledge" like comments and revision history
- Converts to structured Markdown with semantic annotations
- Tracks document metadata (author, revision date, etc.)
"""

from __future__ import annotations

import hashlib
import logging
from pathlib import Path
from typing import Any, Dict, List, Optional

try:
    from docx import Document as DocxDocument
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.text.paragraph import Paragraph
    from docx.table import Table
    from docx.shape import InlineShape
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from src.core.types import Document
from src.libs.loader.base_loader import BaseLoader

logger = logging.getLogger(__name__)


class DocxLoader(BaseLoader):
    """DOCX Loader with semantic-aware parsing.
    
    This loader preserves:
    1. Document structure: headings, lists, tables
    2. Styles: bold, italic, underline, colors
    3. Comments (批注): Extracted as special sections
    4. Revisions (修订): Track changes with author info
    5. Images: Extracted with placeholders
    6. Headers/Footers: Optional inclusion
    
    Configuration:
        extract_comments: Extract comments as separate section (default: True)
        extract_revisions: Extract tracked changes (default: True)
        extract_images: Extract embedded images (default: True)
        include_headers_footers: Include header/footer content (default: True)
        preserve_styles: Preserve text style annotations (default: True)
        image_storage_dir: Base directory for image storage
    """
    
    def __init__(
        self,
        extract_comments: bool = True,
        extract_revisions: bool = True,
        extract_images: bool = True,
        include_headers_footers: bool = True,
        preserve_styles: bool = True,
        image_storage_dir: str | Path = "data/images",
    ):
        """Initialize DOCX Loader.
        
        Args:
            extract_comments: Whether to extract comments.
            extract_revisions: Whether to extract tracked changes.
            extract_images: Whether to extract embedded images.
            include_headers_footers: Whether to include headers/footers.
            preserve_styles: Whether to preserve style annotations.
            image_storage_dir: Directory for storing extracted images.
        """
        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx is required for DocxLoader. "
                "Install with: pip install python-docx"
            )
        
        self.extract_comments = extract_comments
        self.extract_revisions = extract_revisions
        self.extract_images = extract_images
        self.include_headers_footers = include_headers_footers
        self.preserve_styles = preserve_styles
        self.image_storage_dir = Path(image_storage_dir)
    
    def load(self, file_path: str | Path) -> Document:
        """Load and parse a DOCX file.
        
        Args:
            file_path: Path to the DOCX file.
            
        Returns:
            Document with Markdown text and metadata.
            
        Raises:
            FileNotFoundError: If the DOCX file doesn't exist.
            ValueError: If the file is not a valid DOCX.
            RuntimeError: If parsing fails.
        """
        path = self._validate_file(file_path)
        if path.suffix.lower() not in ['.docx', '.doc']:
            raise ValueError(f"File is not a DOCX: {path}")
        
        # Compute document hash
        doc_hash = self._compute_file_hash(path)
        doc_id = f"doc_{doc_hash[:16]}"
        
        # Parse DOCX
        doc = DocxDocument(path)
        
        # Extract metadata from document properties
        metadata = self._extract_metadata(doc, path, doc_hash)
        
        # Extract content sections
        main_content = self._extract_main_content(doc, doc_hash)
        comments_section = self._extract_comments(doc) if self.extract_comments else ""
        revisions_section = self._extract_revisions(doc) if self.extract_revisions else ""
        headers_footers_section = self._extract_headers_footers(doc) if self.include_headers_footers else ""
        
        # Build final markdown content
        sections = []
        
        # Add headers/footers if available
        if headers_footers_section.strip():
            sections.append(headers_footers_section)
        
        # Main content
        sections.append(main_content)
        
        # Add comments section
        if comments_section.strip():
            sections.append(comments_section)
            metadata['has_comments'] = True
        
        # Add revisions section
        if revisions_section.strip():
            sections.append(revisions_section)
            metadata['has_revisions'] = True
        
        # Combine all sections
        text_content = "\n\n".join(sections)
        
        # Extract title
        title = metadata.get('title') or self._extract_title(text_content)
        if title:
            metadata['title'] = title
        
        return Document(
            id=doc_id,
            text=text_content,
            metadata=metadata
        )
    
    def _extract_metadata(
        self, 
        doc: DocxDocument, 
        path: Path, 
        doc_hash: str
    ) -> Dict[str, Any]:
        """Extract document metadata.
        
        Args:
            doc: Parsed DOCX document.
            path: File path.
            doc_hash: Document hash.
            
        Returns:
            Metadata dictionary.
        """
        metadata = {
            "source_path": str(path),
            "doc_type": "docx",
            "doc_hash": doc_hash,
        }
        
        # Core properties
        core_props = doc.core_properties
        if core_props:
            if core_props.title:
                metadata["title"] = core_props.title
            if core_props.author:
                metadata["author"] = core_props.author
            if core_props.created:
                metadata["created_date"] = core_props.created.isoformat()
            if core_props.modified:
                metadata["modified_date"] = core_props.modified.isoformat()
            if core_props.subject:
                metadata["subject"] = core_props.subject
            if core_props.keywords:
                metadata["keywords"] = core_props.keywords
            if core_props.category:
                metadata["category"] = core_props.category
            if core_props.comments:
                metadata["doc_comments"] = core_props.comments
        
        # Document statistics
        metadata["paragraph_count"] = len(doc.paragraphs)
        metadata["table_count"] = len(doc.tables)
        
        return metadata
    
    def _extract_main_content(
        self, 
        doc: DocxDocument, 
        doc_hash: str
    ) -> str:
        """Extract main document content as Markdown.
        
        Args:
            doc: Parsed DOCX document.
            doc_hash: Document hash for image paths.
            
        Returns:
            Markdown-formatted content string.
        """
        markdown_parts = []
        image_index = 0
        
        for element in doc.element.body:
            # Check if paragraph
            if element.tag.endswith('p'):
                para = None
                for p in doc.paragraphs:
                    if p._element == element:
                        para = p
                        break
                
                if para:
                    md_text = self._paragraph_to_markdown(para)
                    if md_text.strip():
                        markdown_parts.append(md_text)
            
            # Check if table
            elif element.tag.endswith('tbl'):
                table = None
                for t in doc.tables:
                    if t._element == element:
                        table = t
                        break
                
                if table:
                    md_table = self._table_to_markdown(table)
                    markdown_parts.append(md_table)
            
            # Check for embedded images
            if self.extract_images:
                for shape in doc.inline_shapes:
                    if shape._element.getparent() == element:
                        image_meta = self._extract_image(shape, doc_hash, image_index)
                        if image_meta:
                            placeholder = f"\n[IMAGE: {image_meta['id']}]\n"
                            markdown_parts.append(placeholder)
                            image_index += 1
        
        return "\n\n".join(markdown_parts)
    
    def _paragraph_to_markdown(self, para: Paragraph) -> str:
        """Convert a paragraph to Markdown.
        
        Args:
            para: DOCX paragraph object.
            
        Returns:
            Markdown-formatted string.
        """
        style_name = para.style.name if para.style else ""
        
        # Get combined text with style markers
        text_parts = []
        for run in para.runs:
            run_text = run.text
            if not run_text:
                continue
            
            # Apply inline formatting
            if self.preserve_styles:
                if run.bold and run.italic:
                    run_text = f"***{run_text}***"
                elif run.bold:
                    run_text = f"**{run_text}**"
                elif run.italic:
                    run_text = f"*{run_text}*"
                elif run.underline:
                    run_text = f"__{run_text}__"
            
            text_parts.append(run_text)
        
        text = "".join(text_parts)
        
        # Determine heading level
        if style_name.startswith('Heading'):
            level = style_name[-1] if style_name[-1].isdigit() else '1'
            prefix = "#" * int(level) + " "
            return f"{prefix}{text}"
        
        # Check for list items
        if para._element.find(qn('w:pPr')) is not None:
            pPr = para._element.find(qn('w:pPr'))
            numPr = pPr.find(qn('w:numPr'))
            if numPr is not None:
                # This is a list item
                return f"- {text}"
        
        # Regular paragraph
        return text
    
    def _table_to_markdown(self, table: Table) -> str:
        """Convert a table to Markdown.
        
        Args:
            table: DOCX table object.
            
        Returns:
            Markdown-formatted table string.
        """
        markdown_lines = []
        
        for row_idx, row in enumerate(table.rows):
            cells = []
            for cell in row.cells:
                cell_text = cell.text.strip().replace('\n', ' ')
                cells.append(cell_text)
            
            line = "| " + " | ".join(cells) + " |"
            markdown_lines.append(line)
            
            # Add header separator after first row
            if row_idx == 0:
                separator = "| " + " | ".join(["---"] * len(cells)) + " |"
                markdown_lines.append(separator)
        
        return "\n".join(markdown_lines)
    
    def _extract_comments(self, doc: DocxDocument) -> str:
        """Extract comments from document.
        
        Args:
            doc: Parsed DOCX document.
            
        Returns:
            Markdown-formatted comments section.
        """
        comments = []
        
        # Access comments via document properties
        # Note: python-docx doesn't directly support comments
        # We need to access the comments.xml part
        try:
            comments_part = doc.part.package.part_related_by.get(
                'http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments'
            )
            if comments_part:
                # Parse comments XML
                from xml.etree import ElementTree as ET
                xml_str = comments_part.blob.decode('utf-8')
                root = ET.fromstring(xml_str)
                
                # Namespace
                ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                
                for comment in root.findall('.//w:comment', ns):
                    comment_id = comment.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id', '')
                    author = comment.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}author', 'Unknown')
                    date = comment.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}date', '')
                    
                    comment_text = ""
                    for p in comment.findall('.//w:p', ns):
                        p_text = ""
                        for r in p.findall('.//w:r', ns):
                            t = r.find('w:t', ns)
                            if t is not None and t.text:
                                p_text += t.text
                        if p_text:
                            comment_text += p_text + "\n"
                    
                    if comment_text.strip():
                        comments.append({
                            'id': comment_id,
                            'author': author,
                            'date': date,
                            'text': comment_text.strip()
                        })
        except Exception as e:
            logger.debug(f"Could not extract comments: {e}")
        
        if not comments:
            return ""
        
        # Format as markdown section
        lines = ["\n\n## 文档批注 (Document Comments)\n"]
        for comment in comments:
            lines.append(f"\n**批注 {comment['id']}** (作者：{comment['author']}, 日期：{comment['date']})")
            lines.append(f"> {comment['text']}")
        
        return "\n".join(lines)
    
    def _extract_revisions(self, doc: DocxDocument) -> str:
        """Extract tracked changes/revisions from document.
        
        Args:
            doc: Parsed DOCX document.
            
        Returns:
            Markdown-formatted revisions section.
        """
        revisions = []
        
        try:
            # Access document XML to find revisions
            from xml.etree import ElementTree as ET
            
            # Find all w:ins (insertions) and w:del (deletions)
            body_xml = doc.element.body.xml
            
            # Parse and find revisions
            root = ET.fromstring(body_xml)
            ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
            
            # Find insertions
            for ins in root.findall('.//w:ins', ns):
                author = ins.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}author', 'Unknown')
                date = ins.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}date', '')
                
                # Get inserted text
                text_parts = []
                for r in ins.findall('.//w:r', ns):
                    t = r.find('w:t', ns)
                    if t is not None and t.text:
                        text_parts.append(t.text)
                
                if text_parts:
                    revisions.append({
                        'type': 'insertion',
                        'author': author,
                        'date': date,
                        'text': ''.join(text_parts)
                    })
            
            # Find deletions
            for del_ in root.findall('.//w:del', ns):
                author = del_.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}author', 'Unknown')
                date = del_.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}date', '')
                
                text_parts = []
                for r in del_.findall('.//w:r', ns):
                    t = r.find('w:t', ns)
                    if t is not None and t.text:
                        text_parts.append(t.text)
                
                if text_parts:
                    revisions.append({
                        'type': 'deletion',
                        'author': author,
                        'date': date,
                        'text': ''.join(text_parts)
                    })
                    
        except Exception as e:
            logger.debug(f"Could not extract revisions: {e}")
        
        if not revisions:
            return ""
        
        # Format as markdown section
        lines = ["\n\n## 修订记录 (Track Changes)\n"]
        for rev in revisions:
            icon = "➕" if rev['type'] == 'insertion' else "➖"
            lines.append(f"\n{icon} **{rev['type'].title()}** by {rev['author']} ({rev['date']})")
            if rev['type'] == 'insertion':
                lines.append(f"> ++{rev['text']}++")
            else:
                lines.append(f"> ~~{rev['text']}~~")
        
        return "\n".join(lines)
    
    def _extract_headers_footers(self, doc: DocxDocument) -> str:
        """Extract headers and footers from document.
        
        Args:
            doc: Parsed DOCX document.
            
        Returns:
            Markdown-formatted headers/footers section.
        """
        headers = []
        footers = []
        
        try:
            # Get sections
            sections = doc.sections
            
            for i, section in enumerate(sections):
                # Header
                if hasattr(section, 'header') and section.header:
                    header_text = section.header.text.strip()
                    if header_text:
                        headers.append(f"(页 {i+1} 页眉) {header_text}")
                
                # Footer
                if hasattr(section, 'footer') and section.footer:
                    footer_text = section.footer.text.strip()
                    if footer_text:
                        footers.append(f"(页 {i+1} 页脚) {footer_text}")
                        
        except Exception as e:
            logger.debug(f"Could not extract headers/footers: {e}")
        
        parts = []
        
        if headers:
            parts.append("\n### 页眉 (Headers)")
            parts.extend(headers)
        
        if footers:
            parts.append("\n### 页脚 (Footers)")
            parts.extend(footers)
        
        return "\n".join(parts) if parts else ""
    
    def _extract_image(
        self, 
        shape: InlineShape, 
        doc_hash: str, 
        index: int
    ) -> Optional[Dict[str, Any]]:
        """Extract an embedded image.
        
        Args:
            shape: DOCX inline shape (image).
            doc_hash: Document hash for path generation.
            index: Image index for unique ID.
            
        Returns:
            Image metadata dictionary or None.
        """
        try:
            # Get image blob
            image_part = shape.image
            image_bytes = image_part.blob
            image_ext = image_part.ext
            
            # Create storage directory
            image_dir = self.image_storage_dir / doc_hash
            image_dir.mkdir(parents=True, exist_ok=True)
            
            # Generate image ID
            image_id = f"{doc_hash[:8]}_img_{index}"
            image_filename = f"{image_id}.{image_ext}"
            image_path = image_dir / image_filename
            
            # Save image
            with open(image_path, 'wb') as f:
                f.write(image_bytes)
            
            # Get dimensions
            width = shape.width.emu if hasattr(shape, 'width') else 0
            height = shape.height.emu if hasattr(shape, 'height') else 0
            
            return {
                'id': image_id,
                'path': str(image_path),
                'dimensions': {
                    'width': width,
                    'height': height
                }
            }
            
        except Exception as e:
            logger.warning(f"Failed to extract image {index}: {e}")
            return None
    
    def _compute_file_hash(self, file_path: Path) -> str:
        """Compute SHA256 hash of file content."""
        sha256 = hashlib.sha256()
        with open(file_path, 'rb') as f:
            for chunk in iter(lambda: f.read(8192), b''):
                sha256.update(chunk)
        return sha256.hexdigest()
    
    def _extract_title(self, text: str) -> Optional[str]:
        """Extract title from first heading or first line."""
        lines = text.split('\n')
        
        # Look for first heading
        for line in lines[:20]:
            line = line.strip()
            if line.startswith('# '):
                return line[2:].strip()
        
        # Fallback to first non-empty line
        for line in lines[:10]:
            line = line.strip()
            if line:
                return line[:100]  # Limit title length
        
        return None